import streamlit as st
import pdfcrowd
from jinja2 import Environment, FileSystemLoader
from docx import Document  # Für .docx Dateien
import re

# PDFCrowd Zugangsdaten
PDFCROWD_USERNAME = 'paulsa'
PDFCROWD_API_KEY = 'e0bd4b588648bfc431efcd2c0df245a2'

st.title("Lebenslauf Generator (Cloud PDF)")

st.header("Kursprofil eingeben")

input_type = st.radio("Kursprofil eingeben oder Datei hochladen?", ("Freitext", "Datei hochladen"))
kursprofil_text = ""

if input_type == "Freitext":
    kursprofil_text = st.text_area("Kursprofil beschreiben", height=200)

elif input_type == "Datei hochladen":
    uploaded_file = st.file_uploader("Lade dein Kursprofil hoch (.txt oder .docx)", type=['txt', 'docx'])
    if uploaded_file:
        if uploaded_file.type == "text/plain":
            kursprofil_text = uploaded_file.read().decode('utf-8')
            st.text_area("Inhalt der Datei", kursprofil_text, height=200)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(uploaded_file)
            kursprofil_text = "\n".join([p.text for p in doc.paragraphs])
            st.text_area("Inhalt der Datei", kursprofil_text, height=200)
        else:
            st.error("Dateityp nicht unterstützt")

# === Bestehenden Lebenslauf hochladen ===
st.header("Bestehenden Lebenslauf hochladen (optional)")

cv_file = st.file_uploader("Bestehenden Lebenslauf hochladen (.txt oder .docx)", type=['txt', 'docx'], key="cv_upload")

parsed_data = {}

if cv_file:
    if cv_file.type == "text/plain":
        cv_text = cv_file.read().decode('utf-8')
    elif cv_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = Document(cv_file)
        cv_text = "\n".join([p.text for p in doc.paragraphs])
    else:
        st.error("Dateityp nicht unterstützt")
        cv_text = ""

    # Einfache Parsing-Logik (Basisversion)
    parsed_data['name'] = re.search(r'Name[:\s]+(.+)', cv_text)
    parsed_data['job_title'] = re.search(r'Berufsbezeichnung[:\s]+(.+)', cv_text)
    parsed_data['birth_year'] = re.search(r'Geburtsjahr[:\s]+(\d{4})', cv_text)
    parsed_data['location'] = re.search(r'Wohnort[:\s]+(.+)', cv_text)
    parsed_data['family_status'] = re.search(r'Familienstand[:\s]+(.+)', cv_text)
    parsed_data['nationality'] = re.search(r'Staatsangehörigkeit[:\s]+(.+)', cv_text)
    parsed_data['career_goal'] = re.search(r'Berufsziel[:\s]+(.+)', cv_text)

    st.success("Bestehender Lebenslauf wurde eingelesen. Die Felder sind vorbefüllt und können bearbeitet werden.")

# Formularfelder (mit Vorbefüllung aus CV oder leer)
st.header("Lebenslauf-Daten eingeben")

job_title = st.text_input("Berufsbezeichnung", parsed_data.get('job_title').group(1) if parsed_data.get('job_title') else "Finanzbuchhalter")
name = st.text_input("Name", parsed_data.get('name').group(1) if parsed_data.get('name') else "Max Mustermann")
birth_year = st.text_input("Geburtsjahr", parsed_data.get('birth_year').group(1) if parsed_data.get('birth_year') else "1990")
location = st.text_input("Wohnort", parsed_data.get('location').group(1) if parsed_data.get('location') else "Berlin")
family_status = st.text_input("Familienstand", parsed_data.get('family_status').group(1) if parsed_data.get('family_status') else "Ledig")
nationality = st.text_input("Staatsangehörigkeit", parsed_data.get('nationality').group(1) if parsed_data.get('nationality') else "Deutsch")
career_goal = st.text_area("Berufsziel", parsed_data.get('career_goal').group(1) if parsed_data.get('career_goal') else "Motivierter Finanzbuchhalter mit Erfahrung...")

education_input = st.text_area("Ausbildung (Einträge mit Semikolon trennen)", "Weiterbildung zum Finanzbuchhalter (IHK); Ausbildung zum Industriekaufmann")
skills_professional = st.text_area("Fachliche Kenntnisse (Semikolon getrennt)", "Buchhaltung; Umsatzsteuer-Voranmeldungen; Monatsabschlüsse")
skills_personal = st.text_area("Persönliche Eigenschaften (Semikolon getrennt)", "Zuverlässig; Teamfähig; Lernbereit")
languages_input = st.text_area("Sprachen (Semikolon getrennt)", "Deutsch; Englisch")

# Fiktive Arbeitserfahrung
work_experience = [
    {
        "position": "Finanzbuchhalter",
        "period": "02/2021 - heute",
        "company": "Dienstleistungsunternehmen (KMU), Berlin",
        "tasks": [
            "Bearbeitung von Eingangs- und Ausgangsrechnungen",
            "Abstimmung von Konten und Unterstützung bei Monatsabschlüssen",
            "Erstellung von Umsatzsteuer-Voranmeldungen",
            "Nutzung von DATEV Unternehmen online und digitalen Buchhaltungslösungen",
            "Zusammenarbeit mit Steuerberatern und internen Abteilungen"
        ]
    },
    {
        "position": "Kaufmännischer Sachbearbeiter",
        "period": "09/2019 - 01/2021",
        "company": "Handelsunternehmen, Berlin",
        "tasks": [
            "Rechnungsprüfung und Bearbeitung des Zahlungsverkehrs",
            "Unterstützung in der vorbereitenden Buchhaltung",
            "Erstellung von Auswertungen in MS Excel"
        ]
    }
]

# Parse education entries
education = []
for edu_str in education_input.split(";"):
    edu_str = edu_str.strip()
    if edu_str:
        education.append({
            "degree": edu_str,
            "period": "",
            "institution": ""
        })

languages = [l.strip() for l in languages_input.split(";") if l.strip()]

skills = {
    "fachkompetenz": [s.strip() for s in skills_professional.split(";") if s.strip()],
    "personal_strengths": [s.strip() for s in skills_personal.split(";") if s.strip()],
    "software": [],
    "languages": languages
}

cv_data = {
    "job_title": job_title,
    "name": name,
    "birth_year": birth_year,
    "location": location,
    "family_status": family_status,
    "nationality": nationality,
    "career_goal": career_goal,
    "work_experience": work_experience,
    "education": education,
    "skills": skills,
    "salary": "",
    "availability": "",
    "email": "",
    "phone": "",
    "kursprofil": kursprofil_text
}

if st.button("Lebenslauf generieren"):
    env = Environment(loader=FileSystemLoader('.'))
    template = env.get_template('cv_template.html')
    html_content = template.render(cv=cv_data)

    try:
        client = pdfcrowd.HtmlToPdfClient(PDFCROWD_USERNAME, PDFCROWD_API_KEY)
        pdf_bytes = client.convertString(html_content)
        output_file = f"{name.replace(' ', '_')}_Lebenslauf.pdf"

        st.success("PDF erfolgreich erstellt!")
        st.download_button("PDF herunterladen", data=pdf_bytes, file_name=output_file, mime="application/pdf")

    except pdfcrowd.Error as e:
        st.error(f"Fehler bei der PDF-Erstellung: {e}")
